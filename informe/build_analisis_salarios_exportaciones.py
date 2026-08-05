# -*- coding: utf-8 -*-
"""Borrador corto: solo salarios + exportaciones (para revisar con el equipo)."""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
PLOTS = ROOT / "outputs" / "plots"
OUT = ROOT / "informe" / "Analisis_indicadores_salarios_exportaciones_borrador_corto.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(t):
    doc.add_heading(t, level=1)


def p(text):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def add_figure(name, caption=None, width_cm=15.5):
    path = PLOTS / name
    if not path.exists():
        p(f"[Gráfico no encontrado: {name}]")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(9)


def ficha(**kwargs):
    p("FICHA METODOLÓGICA")
    p(f"¿Qué mide?: {kwargs['que_mide']}")
    p(f"¿Para qué sirve?: {kwargs['para_que']}")
    p("¿Qué nos dice?:")
    p(f"Si sube: {kwargs['si_sube']}")
    p(f"Si baja: {kwargs['si_baja']}")
    p(f"Comportamiento Esperado: {kwargs['comportamiento']}")
    p(f"Unidad de Medida: {kwargs['unidad']}")
    p(f"Fórmula de Cálculo: {kwargs['formula']}")
    p(f"Fuente de Datos Oficial: {kwargs['fuente']}")
    p(f"Frecuencia de Actualización: {kwargs['frecuencia']}")
    p("Limitaciones:")
    for lim in kwargs["limitaciones"]:
        bullet(lim)


def fmt_money(x):
    return f"${x:,.0f}".replace(",", ".")


# ---- números vivos ----
sipa = pd.read_csv(ROOT / "data/inputs_md/03_salarios_privados_SIPA.csv", parse_dates=["fecha"])
idx_sal = pd.read_csv(
    ROOT / "data/inputs_md/03_salarios_privados_SIPA_indice_region.csv",
    parse_dates=["fecha"],
)
exp = pd.read_csv(ROOT / "data/inputs_md/14_exportaciones_indice_2015_region.csv")

# último mes confiable salarios (evitar provisionales que repiten)
FECHA_SAL = pd.Timestamp("2025-10-01")
noa = ["Catamarca", "Jujuy", "Salta", "Santiago del Estero", "Tucumán"]

sipa_m = sipa[sipa["fecha"] == FECHA_SAL].copy()
lr_nom = float(sipa_m.loc[sipa_m["jurisdiccion"] == "La Rioja", "salario_promedio"].iloc[0])
noa_nom = float(sipa_m.loc[sipa_m["jurisdiccion"].isin(noa), "salario_promedio"].mean())
resto_nom = float(
    sipa_m.loc[
        ~sipa_m["jurisdiccion"].isin(noa + ["La Rioja"]),
        "salario_promedio",
    ].mean()
)
rank = (
    sipa_m.sort_values("salario_promedio", ascending=False)
    .reset_index(drop=True)
)
rank["puesto"] = rank.index + 1
puesto_lr = int(rank.loc[rank["jurisdiccion"] == "La Rioja", "puesto"].iloc[0])

sal_lr = idx_sal[
    (idx_sal["fecha"] == FECHA_SAL) & (idx_sal["la_rioja_region"] == "3. La Rioja")
].iloc[0]
sal_noa = idx_sal[
    (idx_sal["fecha"] == FECHA_SAL) & (idx_sal["la_rioja_region"] == "2. NOA-Resto")
].iloc[0]
sal_resto = idx_sal[
    (idx_sal["fecha"] == FECHA_SAL) & (idx_sal["la_rioja_region"] == "1. Resto país")
].iloc[0]
sal_lr_2017 = idx_sal[
    (idx_sal["fecha"] == "2017-01-01") & (idx_sal["la_rioja_region"] == "3. La Rioja")
].iloc[0]

anio_ult = int(exp["anio"].max())
e_lr = exp[(exp["anio"] == anio_ult) & (exp["la_rioja_region"] == "3. La Rioja")].iloc[0]
e_noa = exp[(exp["anio"] == anio_ult) & (exp["la_rioja_region"] == "2. NOA-Resto")].iloc[0]
e_resto = exp[(exp["anio"] == anio_ult) & (exp["la_rioja_region"] == "1. Resto país")].iloc[0]
e_lr_ant = exp[(exp["anio"] == anio_ult - 1) & (exp["la_rioja_region"] == "3. La Rioja")].iloc[0]
var_exp = 100 * (e_lr["exportaciones_millones_usd"] / e_lr_ant["exportaciones_millones_usd"] - 1)

# =============================================================================
h1("Análisis de indicadores — Salarios y Exportaciones")
p(
    "Versión corta para revisión (estilo empleo/informalidad): pregunta, "
    "metodología breve, último dato, pocos gráficos y ficha. "
    "Salarios actualizados con X-13 + índice real (ene-2025 = 100). "
    "Exportaciones: se mantiene índice 2015 = 100."
)

# =============================================================================
h1("Salarios")
p("¿Cómo se ubican los salarios formales de La Rioja?")

p(
    "Usamos SIPA (OEDE): remuneración promedio del sector privado registrado "
    "por provincia. El nominal sirve para el ranking en un mes. La tendencia "
    "de poder de compra se lee como CEPA/Trabajo: desestacionalizar (X-13), "
    "deflactar por IPC nacional e indexar a ene-2025 = 100. "
    "Por encima de 100 = más poder de compra que ene-2025; por debajo = menos."
)

p("Metodología:")
bullet(
    "Fuente: SIPA / OEDE – boletín provincial de remuneraciones (valores corrientes). "
    "La serie oficial “sin estacionalidad” de Trabajo (trabajoregistrado, hoja A.4) "
    "es solo nacional; acá desestacionalizamos por región (La Rioja / NOA-Resto / Resto país)."
)
bullet(
    "Real = remuneración SA / IPC nacional × 100; después índice con base ene-2025 = 100."
)
bullet(
    "Ventana del gráfico de tendencia: 2017–oct-2025. Los meses posteriores del "
    "Excel provincial repiten valores previos (dato provisorio)."
)

p(f"¿Qué nos dice el último dato nominal (oct-2025)?")
p(
    f"En La Rioja la remuneración promedio del privado registrado fue de "
    f"{fmt_money(lr_nom)}. Eso la ubica en el puesto {puesto_lr} de 24. "
    f"NOA-Resto ≈ {fmt_money(noa_nom)}; resto del país ≈ {fmt_money(resto_nom)}."
)

add_figure(
    "03_salarios_privados_SIPA.png",
    "Figura. Remuneración promedio del sector privado registrado (pesos corrientes).",
)

p("¿Qué nos dice la serie en términos reales?")
p(
    f"Índice real (ene-2025 = 100), oct-2025: La Rioja "
    f"{sal_lr['indice_salario_real_sa']:.0f}; NOA-Resto "
    f"{sal_noa['indice_salario_real_sa']:.0f}; resto país "
    f"{sal_resto['indice_salario_real_sa']:.0f}. "
    f"En ene-2017 La Rioja estaba en "
    f"{sal_lr_2017['indice_salario_real_sa']:.0f}: el poder de compra era mayor "
    f"que en ene-2025. Tras el piso de 2023–24 hubo recuperación; a oct-2025 "
    f"La Rioja queda cerca de la base 100 y por debajo del NOA-Resto."
)

add_figure(
    "03_salarios_privados_SIPA_real_sa_indice.png",
    "Figura. Remuneración privada en términos reales (SA X-13 / IPC), "
    "índice ene-2025 = 100. Estilo CEPA.",
)

ficha(
    que_mide=(
        "Remuneración promedio del empleo asalariado privado registrado (SIPA): "
        "nivel nominal y poder de compra (serie SA deflactada por IPC, índice ene-2025 = 100)."
    ),
    para_que=(
        "Comparar niveles entre jurisdicciones (nominal) y seguir si el salario "
        "privado formal gana o pierde poder de compra (real SA)."
    ),
    si_sube=(
        "En nominal: sube la remuneración corriente. En el índice real: mejora "
        "el poder de compra respecto de ene-2025."
    ),
    si_baja=(
        "En nominal: cae la remuneración corriente. En el índice real: se pierde "
        "poder de compra respecto de ene-2025."
    ),
    comportamiento=(
        "El nominal “por todo concepto” incluye SAC (jun/dic). Por eso se "
        "desestacionaliza con X-13 antes de deflactar."
    ),
    unidad="Pesos corrientes; índice real (ene-2025 = 100).",
    formula=(
        "SA = X-13(nominal_regional); real = SA / IPC_nacional × 100; "
        "índice = 100 × real_t / real_ene2025."
    ),
    fuente="SIPA/OEDE (Capital Humano); IPC INDEC/SSPM; X-13 via paquete seasonal (R).",
    frecuencia="Mensual.",
    limitaciones=[
        "Solo privado registrado (no informal ni monotributo puro).",
        "No separa sector público (para eso EPH).",
        "IPC nacional no captura canasta regional.",
        "Agregado regional = media simple de provincias.",
        "Meses posteriores a oct-2025 en el Excel reciente parecen provisionales.",
    ],
)

# =============================================================================
h1("Exportaciones")
p("¿Cuánto exporta La Rioja?")

p(
    "Fuente: OPEX–INDEC (millones de USD por origen provincial). "
    "Se agregan La Rioja, NOA-Resto y Resto país por suma. "
    "El índice 2015 = 100 se mantiene (no es el mismo caso que salarios en pesos; "
    "confirmado con el equipo)."
)

p("Metodología:")
bullet(
    "Nivel: millones de USD (paneles con ejes libres). Dinámica: índice 2015 = 100. "
    "Peso: % sobre la suma OPEX de las 24 jurisdicciones."
)
bullet(
    "El share usa el total por origen provincial, no la suma de subrubros "
    "detallados (OPEX no publica todos los productos todos los años)."
)
bullet("2024–2025 pueden figurar como provisorios.")

p(f"¿Qué nos dice el último dato ({anio_ult})?")
bullet(
    f"La Rioja exportó {e_lr['exportaciones_millones_usd']:.0f} millones de USD "
    f"({e_lr['share_nacional_pct']:.2f}% de la suma provincial OPEX)."
)
bullet(
    f"Respecto de {anio_ult - 1} ({e_lr_ant['exportaciones_millones_usd']:.0f} millones): "
    f"{'cayeron' if var_exp < 0 else 'crecieron'} {abs(var_exp):.0f}% "
    f"(dato de La Rioja)."
)
bullet(
    f"Índice 2015 = 100: La Rioja {e_lr['indice_2015']:.0f}; "
    f"NOA-Resto {e_noa['indice_2015']:.0f}; resto país {e_resto['indice_2015']:.0f}."
)

add_figure(
    "14_exportaciones_totales_region.png",
    "Figura. Exportaciones en millones de USD (ejes libres por región).",
)
add_figure(
    "14_exportaciones_indice_2015_region.png",
    "Figura. Índice 2015 = 100 (compara dinámicas, no niveles absolutos).",
)
add_figure(
    "14_exportaciones_share_nacional_region.png",
    "Figura. Participación sobre la suma OPEX de las 24 jurisdicciones.",
)

p("¿En qué se concentra lo que exporta La Rioja?")
p(
    "La Rioja está más concentrada que el NOA-Resto y el resto del país: "
    "el top 3 concentra ~77% de lo publicado (preparados hortifrutícolas, "
    "grasas/aceites, papel/cartón), frente a ~35% en los otros agregados. "
    "Eso eleva la exposición a shocks de esos eslabones."
)
p(
    "El perfil también es distinto. El resto del país se parece al patrón "
    "agroexportador/energético (cereales, residuos alimentarios, grasas y "
    "aceites, transporte, petróleo). El NOA-Resto combina químicos, "
    "cereales, oleaginosos y minería, más hortalizas y tabaco. La Rioja se "
    "apoya en agroindustria elaborada, aceites y papel/cartón; casi no "
    "aparece el núcleo cerealero ni el minero-energético. Que OPEX no "
    "publique un bloque minero desagregado no prueba cero minería (puede "
    "ser residual o confidencialidad)."
)

add_figure(
    "14_exportaciones_treemap_subrubros_facet.png",
    "Figura. Composición por subrubros (2025).",
    width_cm=16,
)
add_figure(
    "14_exportaciones_heatmap_larioja.png",
    "Figura. Heatmap de subrubros — La Rioja.",
)

ficha(
    que_mide=(
        "Exportaciones de bienes por origen provincial/grupo (millones USD); "
        "índice 2015 = 100; share sobre la suma de totales provinciales OPEX."
    ),
    para_que=(
        "Seguir inserción externa, comparar dinámicas sin confundir escalas y "
        "ver especialización por productos."
    ),
    si_sube="Mayor venta al exterior o mejores precios/cantidades.",
    si_baja="Menor inserción externa o caída en rubros clave.",
    comportamiento=(
        "Cíclico; sensible a precios internacionales, tipo de cambio y producción local."
    ),
    unidad="Millones de USD; índice 2015 = 100; % sobre suma OPEX provincial.",
    formula=(
        "Total grupo = suma de provincias. "
        "Índice_t = 100 × exp_t / exp_2015. "
        "Share = exp_grupo / suma_24_jurisdicciones × 100."
    ),
    fuente="INDEC – OPEX.",
    frecuencia="Anual (con revisiones / datos provisorios).",
    limitaciones=[
        "No publica todos los productos todos los años.",
        "No incluye servicios.",
        "La suma de subrubros no cierra con el total provincial.",
        "El índice compara dinámica, no tamaño absoluto.",
    ],
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print("OK", OUT)
